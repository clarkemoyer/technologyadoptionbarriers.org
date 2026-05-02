#!/usr/bin/env python3
"""
Merge 4 consolidated appendix markdown files (A-D) into the CRP Word document.
Appendix A: TABS Survey Instrument (concept mapping, V2/V1 instruments, revision record, live ops)
Appendix B: TABS Research Platform (GitHub, blog, bibliography, personas, AI dev, integrations)
Appendix C: Data Analysis and Validation (comprehensive report, scripts, validation registry)
Appendix D: Institutional Governance (doctoral committee, IRB)
Auto-discovers latest CRP body DOCX and appendix files. Handles landscape orientation.

Usage:
    python merge_appendixes.py
    python merge_appendixes.py --workspace /path/to/CRP-workspace
    python merge_appendixes.py --docx /path/to/body.docx --appendix-dir /path/to/appendixes
    python merge_appendixes.py --docx /path/to/body.docx --appendix-dir /path/to/appendixes \\
        --output /path/to/output.docx
"""

import argparse
import glob as globmod
import os
import re
import sys
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, Twips

# Shared portable workspace discovery (replaces the old hardcoded
# /sessions/*/... globs). See scripts/crp-document-tools/paths.py for
# the full discovery precedence and the Quickstart-for-Researchers tour.
_SUBTREE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _SUBTREE_DIR not in sys.path:
    sys.path.insert(0, _SUBTREE_DIR)
from paths import (  # noqa: E402
    CrpWorkspaceNotFound,
    find_workspace as _find_workspace_shared,
    parse_filename_date,
)

# ── Configuration constants ─────────────────────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(9)
SMALL_TABLE_FONT_SIZE = Pt(8)
LINE_SPACING = Pt(24)  # double spacing for 12pt

# Appendixes that need landscape orientation for their tables
LANDSCAPE_APPENDIXES = {'A', 'B'}

# Consolidated 4-appendix structure (A-D), each file auto-discovered by prefix
APPENDIX_FILES = [
    ("A", "TABS Survey Instrument"),
    ("B", "TABS Research Platform"),
    ("C", "Data Analysis and Validation"),
    ("D", "Institutional Governance"),
]

# ── Workspace / path discovery ──────────────────────────────────────────

def find_workspace():
    """Locate the CRP workspace root via the shared portable discovery.

    Returns the workspace path, or None if discovery fails. (We swallow
    the exception so callers can supply --docx + --appendix-dir and skip
    workspace discovery entirely.)
    """
    try:
        return _find_workspace_shared()
    except CrpWorkspaceNotFound:
        return None


def find_latest_docx(workspace):
    """Find the latest CRP body .docx in '01 CRP Body/' of the workspace.

    Uses date-aware sorting on the (M-D-YYYY HHMM TZ) timestamp embedded
    in the filename (DST-tolerant via paths.parse_filename_date), falling
    back to file mtime when the filename has no embedded timestamp.
    """
    body_dir = os.path.join(workspace, "01 CRP Body")
    if not os.path.isdir(body_dir):
        return None
    candidates = []
    for f in os.listdir(body_dir):
        if f.startswith("Clarke Moyer") and f.endswith(".docx") and not f.startswith("~"):
            candidates.append(os.path.join(body_dir, f))
    if not candidates:
        for f in os.listdir(body_dir):
            if f.endswith(".docx") and not f.startswith("~"):
                candidates.append(os.path.join(body_dir, f))
    if not candidates:
        return None

    def _sort_key(p):
        parsed = parse_filename_date(os.path.basename(p))
        # If filename has no embedded timestamp (year=0), fall back to mtime
        if parsed == (0, 0, 0, 0):
            try:
                return (0, 0, 0, int(os.path.getmtime(p)))
            except OSError:
                return (0, 0, 0, 0)
        return parsed

    return max(candidates, key=_sort_key)


def _find_appendix(prefix, appendix_dir):
    """Find the latest appendix file matching a prefix, using date-aware sorting."""
    candidates = globmod.glob(os.path.join(appendix_dir, f"{prefix}*.md"))
    if not candidates:
        return None
    return max(candidates, key=lambda f: parse_filename_date(os.path.basename(f)))


# ── Helper Functions ───────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=None):
    """Set font properties on a run."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    # Set east-asian and complex script fonts too
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=480):
    """Set paragraph spacing. line_spacing in twips (480 = double for 12pt)."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        paragraph._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}" w:line="{line_spacing}" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'), str(after))
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'auto')


def add_body_paragraph(doc, text, first_line_indent=True, bold=False, italic=False, alignment=None, spacing_after=0):
    """Add a standard body paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    set_paragraph_spacing(p, before=0, after=spacing_after, line_spacing=480)

    if first_line_indent:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            p._element.insert(0, pPr)
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="720"/>')
        pPr.append(ind)

    add_formatted_runs(p, text, bold=bold, italic=italic)
    return p


def add_formatted_runs(paragraph, text, bold=False, italic=False, font_size=None):
    """Add runs with inline formatting (bold, italic) from markdown text."""
    # Handle inline formatting: **bold**, *italic*, ***bold+italic***, `code`
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

    parts = re.split(pattern, text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if part is None:
            i += 1
            continue
        # Check if this is a match group
        if i + 5 <= len(parts) and parts[i] and re.match(r'^\*{1,3}.*\*{1,3}$|^`.*`$', parts[i]):
            # This is the full match, skip - use the capture groups
            if parts[i+1]:  # ***bold+italic***
                run = paragraph.add_run(parts[i+1])
                set_run_font(run, bold=True, italic=True, size=font_size)
            elif parts[i+2]:  # **bold**
                run = paragraph.add_run(parts[i+2])
                set_run_font(run, bold=True, italic=False, size=font_size)
            elif parts[i+3]:  # *italic*
                run = paragraph.add_run(parts[i+3])
                set_run_font(run, bold=False, italic=True, size=font_size)
            elif parts[i+4]:  # `code`
                run = paragraph.add_run(parts[i+4])
                set_run_font(run, bold=False, italic=False, size=font_size)
                run.font.name = "Courier New"
            i += 5
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, bold=bold, italic=italic, size=font_size)
            i += 1


def add_heading(doc, text, level, centered=False):
    """Add a heading paragraph (appears in TOC)."""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=240 if level > 1 else 0, after=120, line_spacing=480)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_appendix_subheading(doc, text, level, centered=False):
    """Add a heading-styled paragraph that does NOT appear in the TOC.
    Uses Normal style with bold formatting and appropriate sizing to
    visually match Heading 2/3 without triggering TOC inclusion."""
    p = doc.add_paragraph()
    # Use Normal style — NOT Heading 2/3 — so TOC ignores it
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=240, after=120, line_spacing=480)
    run = p.add_run(text)
    # Match heading appearance: bold, slightly larger for H2 vs H3
    if level == 2:
        set_run_font(run, bold=True, size=Pt(13))
    else:
        set_run_font(run, bold=True, size=Pt(12))
    return p


def add_appendix_title_page(doc, letter, subtitle):
    """Add an appendix title: 'APPENDIX X: Subtitle' centered as Heading1."""
    # Page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    # Combined title - "APPENDIX X: Subtitle" as single Heading1 for TOC navigation
    title_p = doc.add_paragraph()
    title_p.style = doc.styles['Heading 1']
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_p, after=240, line_spacing=480)
    run = title_p.add_run(f"APPENDIX {letter}: {subtitle}")
    set_run_font(run, bold=True)

    return title_p


def parse_markdown_table(lines):
    """Parse markdown table lines into rows of cells. Returns (headers, rows)."""
    if len(lines) < 2:
        return [], []

    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = parse_row(lines[0])
    # Skip separator line (lines[1] which is |---|---|...)
    rows = []
    for line in lines[2:]:
        if line.strip():
            rows.append(parse_row(line))

    return headers, rows


def add_table(doc, headers, rows, landscape=False, small_font=False):
    """Add a formatted table to the document."""
    if not headers:
        return

    num_cols = len(headers)

    # Calculate column widths
    if landscape:
        content_width = 12960  # 15840 - 2*1440 in DXA (landscape with 1" margins)
    else:
        content_width = 9360   # 12240 - 2*1440 in DXA (portrait with 1" margins)

    # Equal column widths as default
    col_width = content_width // num_cols
    col_widths = [col_width] * num_cols
    # Adjust last column to absorb rounding
    col_widths[-1] = content_width - col_width * (num_cols - 1)

    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{content_width}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(content_width))
        tblW.set(qn('w:type'), 'dxa')

    fs = SMALL_TABLE_FONT_SIZE if small_font else TABLE_FONT_SIZE

    # Style header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        # Set cell width
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tcPr)
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[i]}" w:type="dxa"/>')
        tcPr.append(tcW)
        # Header shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5E8F0" w:val="clear"/>')
        tcPr.append(shading)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=40, after=40, line_spacing=240)
        add_formatted_runs(p, header_text, bold=True, font_size=fs)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            # Set cell width
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[col_idx]}" w:type="dxa"/>')
            tcPr.append(tcW)

            # Alternating row shading
            if row_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
                tcPr.append(shading)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=20, after=20, line_spacing=240)
            add_formatted_runs(p, row_data[col_idx], font_size=fs)

    # Add table borders
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

    # Cell margins
    cell_margins = parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
        <w:top w:w="40" w:type="dxa"/>
        <w:left w:w="80" w:type="dxa"/>
        <w:bottom w:w="40" w:type="dxa"/>
        <w:right w:w="80" w:type="dxa"/>
    </w:tblCellMar>''')
    tblPr.append(cell_margins)

    return table


def add_section_break(doc, landscape=False):
    """Add a section break, optionally switching to/from landscape."""
    new_section = doc.add_section()
    if landscape:
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Twips(15840)
        new_section.page_height = Twips(12240)
    else:
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Twips(12240)
        new_section.page_height = Twips(15840)
    new_section.top_margin = Twips(1440)
    new_section.bottom_margin = Twips(1440)
    new_section.left_margin = Twips(1440)
    new_section.right_margin = Twips(1440)
    return new_section


def process_markdown_content(doc, lines, landscape=False):
    """Process markdown lines and add them to the document."""
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip the H1 title (# Appendix X: ...) - we already added the title page
        if line.startswith('# ') and i < 5:
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            # Add some spacing instead of a visible rule
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=120, after=120, line_spacing=240)
            i += 1
            continue

        # End of appendix marker
        if line.strip().startswith('*End of Appendix') or line.strip().startswith('_End of Appendix'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=240, after=0, line_spacing=480)
            clean_text = line.strip().strip('*').strip('_').strip()
            run = p.add_run(clean_text)
            set_run_font(run, italic=True)
            i += 1
            continue

        # Headings
        if line.startswith('####'):
            text = line.lstrip('#').strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=200, after=100, line_spacing=480)
            add_formatted_runs(p, text, bold=True)
            i += 1
            continue

        if line.startswith('###'):
            text = line.lstrip('#').strip()
            h = add_appendix_subheading(doc, text, 3, centered=False)
            i += 1
            continue

        if line.startswith('##'):
            text = line.lstrip('#').strip()
            h = add_heading(doc, text, 2, centered=False)
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j].strip():
                table_lines.append(lines[j])
                j += 1

            headers, rows = parse_markdown_table(table_lines)
            if headers:
                # Determine if this needs small font
                num_cols = len(headers)
                small = num_cols >= 6 or (landscape and num_cols >= 4)
                add_table(doc, headers, rows, landscape=landscape, small_font=small)
                # Add spacing after table
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=120, after=120, line_spacing=240)

            i = j
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                text = re.sub(r'^\d+\.\s*', '', lines[i].strip())
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=60, line_spacing=480)
                # Add indent
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
                pPr.append(ind)

                # Extract number
                num_match = re.match(r'^(\d+)\.\s', lines[i].strip())
                num = num_match.group(1) if num_match else "1"

                run = p.add_run(f"{num}. ")
                set_run_font(run, bold=False)
                add_formatted_runs(p, text)
                i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                text = lines[i].strip()[2:]

                # Check for sub-items (indented bullets)
                indent_level = len(lines[i]) - len(lines[i].lstrip())

                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=60, line_spacing=480)

                left_indent = 720 if indent_level < 4 else 1440
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_indent}" w:hanging="360"/>')
                pPr.append(ind)

                # Use actual bullet character
                bullet_char = "\u2022"
                run = p.add_run(f"{bullet_char} ")
                set_run_font(run)
                add_formatted_runs(p, text)
                i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            i += 1  # skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```

            # Add code block as formatted paragraph
            for code_line in code_lines:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=0, line_spacing=240)
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
                pPr.append(ind)
                # Add shading
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                pPr.append(shd)

                run = p.add_run(code_line.rstrip())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            # Check if it's a continuation of multi-line content
            full_text = text
            # Don't merge lines that start special patterns
            while (i + 1 < len(lines) and
                   lines[i+1].strip() and
                   not lines[i+1].strip().startswith('#') and
                   not lines[i+1].strip().startswith('-') and
                   not lines[i+1].strip().startswith('*') and
                   not lines[i+1].strip().startswith('|') and
                   not lines[i+1].strip().startswith('```') and
                   not lines[i+1].strip() in ('---', '***', '___') and
                   not re.match(r'^\d+\.\s', lines[i+1].strip()) and
                   not lines[i+1].strip().startswith('*End of') and
                   not lines[i+1].strip().startswith('_End of')):
                i += 1
                full_text += ' ' + lines[i].strip()

            add_body_paragraph(doc, full_text, first_line_indent=True)

        i += 1


# ── Main Execution ─────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Merge CRP body + appendixes into final docx")
    parser.add_argument("--workspace", help="Path to CRP workspace root (auto-discovered if omitted)")
    parser.add_argument("--docx", help="Path to CRP body .docx (overrides workspace discovery)")
    parser.add_argument("--appendix-dir", help="Path to '02 CRP Appendixes' folder (overrides workspace discovery)")
    parser.add_argument("--output", help="Output .docx path (default: timestamped file in workspace root)")
    args = parser.parse_args()

    # Resolve workspace
    workspace = args.workspace or find_workspace()

    # Resolve CRP body docx
    crp_path = args.docx
    if crp_path is None and workspace:
        crp_path = find_latest_docx(workspace)
    if crp_path is None:
        print("ERROR: No CRP body DOCX found. Provide --workspace or --docx.")
        sys.exit(1)

    # Resolve appendix directory
    appendix_dir = args.appendix_dir
    if appendix_dir is None and workspace:
        appendix_dir = os.path.join(workspace, "02 CRP Appendixes")
    if appendix_dir is None:
        print("ERROR: No appendix directory found. Provide --workspace or --appendix-dir.")
        sys.exit(1)

    # Resolve output path
    if args.output:
        output_path = args.output
    else:
        _est = ZoneInfo('America/New_York')
        _now = datetime.now(_est)
        _ts = f"{_now.month}-{_now.day}-{_now.year} {_now.strftime('%H%M')} {_now.strftime('%Z')}"
        out_root = workspace if workspace else os.path.dirname(crp_path)
        output_path = os.path.join(out_root, f"Clarke Moyer - DBA Culminating Research Project - Product Development ({_ts}).docx")

    print(f"Opening CRP body: {os.path.basename(crp_path)}")
    doc = Document(crp_path)

    # Discover appendix files
    appendix_files = []
    for letter, title in APPENDIX_FILES:
        filepath = _find_appendix(f"Appendix {letter}", appendix_dir)
        if filepath is None:
            print(f"ERROR: No file found for Appendix {letter} in {appendix_dir}")
            sys.exit(1)
        appendix_files.append((letter, filepath, title))
        print(f"  Found Appendix {letter}: {os.path.basename(filepath)}")

    # Find the "List of Appendices" heading in the body content section
    # (not in TOC - we need the one that appears as actual content)
    body = doc.element.body
    paragraphs = body.findall(qn('w:p'))

    # Find all paragraphs and identify where to cut
    list_of_appendices_idx = None
    for idx, p_elem in enumerate(paragraphs):
        text_parts = []
        for t in p_elem.itertext():
            text_parts.append(t)
        full_text = ''.join(text_parts).strip()

        # Look for the "List of Appendices" that appears in the actual body content
        # (after the TOC). We want the LAST occurrence.
        if 'List of Appendices' in full_text or 'LIST OF APPENDICES' in full_text:
            # Check if this uses Heading1 style (content section, not TOC entry)
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                    list_of_appendices_idx = idx

    if list_of_appendices_idx is None:
        print("ERROR: Could not find 'List of Appendices' heading in document body!")
        sys.exit(1)

    print(f"Found 'List of Appendices' at paragraph index {list_of_appendices_idx}")
    print(f"Total paragraphs: {len(paragraphs)}")

    # Remove everything from the "List of Appendices" paragraph onward
    # Also remove tables and section properties that come after
    elements_to_remove = []
    found_start = False
    for child in list(body):
        if child.tag == qn('w:p') and child is paragraphs[list_of_appendices_idx]:
            found_start = True
        if found_start and child.tag != qn('w:sectPr'):
            elements_to_remove.append(child)

    print(f"Removing {len(elements_to_remove)} elements (old appendix content)...")
    for elem in elements_to_remove:
        body.remove(elem)

    print("Adding new List of Appendices...")

    # ── Add "List of Appendices" heading ──
    h1 = doc.add_paragraph()
    h1.style = doc.styles['Heading 1']
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(h1, before=0, after=240, line_spacing=480)
    run = h1.add_run("List of Appendices")
    set_run_font(run, bold=True)

    # Add entries for each appendix
    for letter, filepath, title in appendix_files:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line_spacing=480)
        run = p.add_run(f"Appendix {letter}: {title}")
        set_run_font(run)

    # Add blank paragraph after list
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=480)

    # ── Process each appendix ──
    current_landscape = False

    for letter, filepath, title in appendix_files:
        print(f"Processing Appendix {letter}: {title}...")

        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        needs_landscape = letter in LANDSCAPE_APPENDIXES

        # If switching orientation, add a section break
        if needs_landscape and not current_landscape:
            add_section_break(doc, landscape=True)
            current_landscape = True
        elif not needs_landscape and current_landscape:
            add_section_break(doc, landscape=False)
            current_landscape = False

        # Add title page
        add_appendix_title_page(doc, letter, title)

        # Process content
        process_markdown_content(doc, lines, landscape=needs_landscape)

    # If we ended in landscape, switch back to portrait for the final section
    if current_landscape:
        add_section_break(doc, landscape=False)

    # ── Update the final section properties ──
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        pgSz = sect_pr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.set(qn('w:w'), '12240')
            pgSz.set(qn('w:h'), '15840')
            if pgSz.get(qn('w:orient')):
                del pgSz.attrib[qn('w:orient')]

    # Save
    print(f"Saving to: {output_path}")
    doc.save(output_path)
    print("Done!")
    print(f"File size: {os.path.getsize(output_path):,} bytes")


if __name__ == '__main__':
    main()
