#!/usr/bin/env python3
"""
Merge 4 markdown appendix files (A-D) into the CRP Word document.
  A: TABS Survey Instrument
  B: TABS Research Platform
  C: Data Analysis and Validation
  D: Institutional Governance

Uses the root CRP as the base document, strips old appendix content
from the first APPENDIX heading onward, then appends new A-D content
from markdown sources. Preserves the existing List of Appendices
(already correct for A-D in the root CRP).

ARCHIVE NOTE: This is an archived variant (v2) kept for reference only.
It hard-codes session-specific absolute paths and timestamped filenames and
is intentionally non-portable - it will not run outside the original author
workspace. Use the current merge_appendixes.py (in the parent folder) for
all new work; it uses glob-based workspace discovery and accepts CLI args.
"""

import re
import os
import sys
import docx.enum.text
from docx import Document
from docx.shared import Pt, Inches, Twips, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# -- Configuration --

CRP_ROOT = "/sessions/wonderful-focused-rubin/mnt/! Clarke Moyer Smeal CRP - TABS"

CRP_PATH = os.path.join(CRP_ROOT, "Clarke Moyer - DBA Culminating Research Project - Product Development (4-11-2026 2105 EST).docx")
APPENDIX_DIR = os.path.join(CRP_ROOT, "02 CRP Appendixes")
OUTPUT_PATH = os.path.join(CRP_ROOT, "Clarke Moyer - DBA Culminating Research Project - Product Development (4-16-2026 0432 EST).docx")

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
TABLE_FONT_SIZE = Pt(9)
SMALL_TABLE_FONT_SIZE = Pt(8)
LINE_SPACING = Pt(24)  # double spacing for 12pt

# Appendixes that need landscape orientation for their tables
LANDSCAPE_APPENDIXES = {'A', 'B', 'C'}

APPENDIX_FILES = [
    ("A", "Appendix A - TABS Survey Instrument (4-15-2026 2253 EST).md", "TABS Survey Instrument"),
    ("B", "Appendix B - TABS Research Platform (4-15-2026 1823 EST).md", "TABS Research Platform"),
    ("C", "Appendix C - Data Analysis and Validation (4-16-2026 0235 EST).md", "Data Analysis and Validation"),
    ("D", "Appendix D - Institutional Governance (4-11-2026 1553 EST).md", "Institutional Governance"),
]


# -- Helper Functions --

def set_run_font(run, bold=False, italic=False, size=None):
    """Set font properties on a run."""
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=480):
    """Set paragraph spacing. line_spacing in twips (480 = double for 12pt)."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        paragraph._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}" w:line="{line_spacing}" w:lineRule="auto"/>')
        pPr.append(spacing)
    else:
        spacing.set(qn('w:before'), str(before))
        spacing.set(qn('w:after'), str(after))
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'auto')


def add_body_paragraph(doc, text, first_line_indent=True, bold=False, italic=False, alignment=None, spacing_after=0):
    """Add a standard body paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    set_paragraph_spacing(p, before=0, after=spacing_after, line_spacing=480)

    if first_line_indent:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            p._element.insert(0, pPr)
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:firstLine="720"/>')
        pPr.append(ind)

    add_formatted_runs(p, text, bold=bold, italic=italic)
    return p


def add_formatted_runs(paragraph, text, bold=False, italic=False, font_size=None):
    """Add runs with inline formatting (bold, italic) from markdown text."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'

    parts = re.split(pattern, text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if part is None:
            i += 1
            continue
        if i + 5 <= len(parts) and parts[i] and re.match(r'^\*{1,3}.*\*{1,3}$|^`.*`$', parts[i]):
            if parts[i+1]:  # ***bold+italic***
                run = paragraph.add_run(parts[i+1])
                set_run_font(run, bold=True, italic=True, size=font_size)
            elif parts[i+2]:  # **bold**
                run = paragraph.add_run(parts[i+2])
                set_run_font(run, bold=True, italic=False, size=font_size)
            elif parts[i+3]:  # *italic*
                run = paragraph.add_run(parts[i+3])
                set_run_font(run, bold=False, italic=True, size=font_size)
            elif parts[i+4]:  # `code`
                run = paragraph.add_run(parts[i+4])
                set_run_font(run, bold=False, italic=False, size=font_size)
                run.font.name = "Courier New"
            i += 5
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, bold=bold, italic=italic, size=font_size)
            i += 1


def add_heading(doc, text, level, centered=False):
    """Add a heading paragraph (appears in TOC)."""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=240 if level > 1 else 0, after=120, line_spacing=480)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_appendix_subheading(doc, text, level, centered=False):
    """Add a heading-styled paragraph that does NOT appear in the TOC.
    Uses Normal style with bold formatting and appropriate sizing."""
    p = doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=240, after=120, line_spacing=480)
    run = p.add_run(text)
    if level == 2:
        set_run_font(run, bold=True, size=Pt(13))
    else:
        set_run_font(run, bold=True, size=Pt(12))
    return p


def add_appendix_title_page(doc, letter, subtitle):
    """Add an appendix title: 'APPENDIX X: Subtitle' centered as Heading1."""
    # Page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

    # Combined title as single Heading1 for TOC navigation
    title_p = doc.add_paragraph()
    title_p.style = doc.styles['Heading 1']
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_p, after=240, line_spacing=480)
    run = title_p.add_run(f"APPENDIX {letter}: {subtitle}")
    set_run_font(run, bold=True)

    return title_p


def parse_markdown_table(lines):
    """Parse markdown table lines into rows of cells. Returns (headers, rows)."""
    if len(lines) < 2:
        return [], []

    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = parse_row(lines[0])
    rows = []
    for line in lines[2:]:
        if line.strip():
            rows.append(parse_row(line))

    return headers, rows


def add_table(doc, headers, rows, landscape=False, small_font=False):
    """Add a formatted table to the document."""
    if not headers:
        return

    num_cols = len(headers)

    if landscape:
        content_width = 12960  # 15840 - 2*1440 in DXA (landscape with 1" margins)
    else:
        content_width = 9360   # 12240 - 2*1440 in DXA (portrait with 1" margins)

    col_width = content_width // num_cols
    col_widths = [col_width] * num_cols
    col_widths[-1] = content_width - col_width * (num_cols - 1)

    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{content_width}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(content_width))
        tblW.set(qn('w:type'), 'dxa')

    fs = SMALL_TABLE_FONT_SIZE if small_font else TABLE_FONT_SIZE

    # Style header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            tc.insert(0, tcPr)
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[i]}" w:type="dxa"/>')
        tcPr.append(tcW)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5E8F0" w:val="clear"/>')
        tcPr.append(shading)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=40, after=40, line_spacing=240)
        add_formatted_runs(p, header_text, bold=True, font_size=fs)

    # Style data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx in range(min(len(row_data), num_cols)):
            cell = table.rows[row_idx + 1].cells[col_idx]
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[col_idx]}" w:type="dxa"/>')
            tcPr.append(tcW)

            if row_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
                tcPr.append(shading)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=20, after=20, line_spacing=240)
            add_formatted_runs(p, row_data[col_idx], font_size=fs)

    # Add table borders
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

    cell_margins = parse_xml(f'''<w:tblCellMar {nsdecls("w")}>
        <w:top w:w="40" w:type="dxa"/>
        <w:left w:w="80" w:type="dxa"/>
        <w:bottom w:w="40" w:type="dxa"/>
        <w:right w:w="80" w:type="dxa"/>
    </w:tblCellMar>''')
    tblPr.append(cell_margins)

    return table


def add_section_break(doc, landscape=False):
    """Add a section break, optionally switching to/from landscape."""
    new_section = doc.add_section()
    if landscape:
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Twips(15840)
        new_section.page_height = Twips(12240)
    else:
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Twips(12240)
        new_section.page_height = Twips(15840)
    new_section.top_margin = Twips(1440)
    new_section.bottom_margin = Twips(1440)
    new_section.left_margin = Twips(1440)
    new_section.right_margin = Twips(1440)
    return new_section


def process_markdown_content(doc, lines, landscape=False):
    """Process markdown lines and add them to the document."""
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Skip the H1 title (# Appendix X: ...) - we already added the title page
        if line.startswith('# ') and i < 5:
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=120, after=120, line_spacing=240)
            i += 1
            continue

        # End of appendix marker
        if line.strip().startswith('*End of Appendix') or line.strip().startswith('_End of Appendix'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=240, after=0, line_spacing=480)
            clean_text = line.strip().strip('*').strip('_').strip()
            run = p.add_run(clean_text)
            set_run_font(run, italic=True)
            i += 1
            continue

        # Headings (#### before ### before ##)
        if line.startswith('####'):
            text = line.lstrip('#').strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=200, after=100, line_spacing=480)
            add_formatted_runs(p, text, bold=True)
            i += 1
            continue

        if line.startswith('###'):
            text = line.lstrip('#').strip()
            add_appendix_subheading(doc, text, 3, centered=False)
            i += 1
            continue

        if line.startswith('##'):
            text = line.lstrip('#').strip()
            add_heading(doc, text, 2, centered=False)
            i += 1
            continue

        # Blockquote (used for scope summaries in the appendixes)
        if line.strip().startswith('>'):
            quote_text = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=60, after=60, line_spacing=480)
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                p._element.insert(0, pPr)
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720"/>')
            pPr.append(ind)
            add_formatted_runs(p, quote_text, italic=True)
            i += 1
            continue

        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j].strip():
                table_lines.append(lines[j])
                j += 1

            headers, rows = parse_markdown_table(table_lines)
            if headers:
                num_cols = len(headers)
                small = num_cols >= 6 or (landscape and num_cols >= 4)
                add_table(doc, headers, rows, landscape=landscape, small_font=small)
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=120, after=120, line_spacing=240)

            i = j
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                text = re.sub(r'^\d+\.\s*', '', lines[i].strip())
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=60, line_spacing=480)
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
                pPr.append(ind)

                num_match = re.match(r'^(\d+)\.\s', lines[i].strip())
                num = num_match.group(1) if num_match else "1"

                run = p.add_run(f"{num}. ")
                set_run_font(run, bold=False)
                add_formatted_runs(p, text)
                i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                text = lines[i].strip()
                if text.startswith('- '):
                    text = text[2:]
                elif text.startswith('* '):
                    text = text[2:]

                indent_level = len(lines[i]) - len(lines[i].lstrip())

                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=60, line_spacing=480)

                left_indent = 720 if indent_level < 4 else 1440
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="{left_indent}" w:hanging="360"/>')
                pPr.append(ind)

                bullet_char = "\u2022"
                run = p.add_run(f"{bullet_char} ")
                set_run_font(run)
                add_formatted_runs(p, text)
                i += 1
            continue

        # Code block
        if line.strip().startswith('```'):
            i += 1  # skip opening ```
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```

            for code_line in code_lines:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=0, line_spacing=240)
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                    p._element.insert(0, pPr)
                ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
                pPr.append(ind)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                pPr.append(shd)

                run = p.add_run(code_line.rstrip())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            full_text = text
            while (i + 1 < len(lines) and
                   lines[i+1].strip() and
                   not lines[i+1].strip().startswith('#') and
                   not lines[i+1].strip().startswith('-') and
                   not lines[i+1].strip().startswith('*') and
                   not lines[i+1].strip().startswith('>') and
                   not lines[i+1].strip().startswith('|') and
                   not lines[i+1].strip().startswith('```') and
                   not lines[i+1].strip() in ('---', '***', '___') and
                   not re.match(r'^\d+\.\s', lines[i+1].strip()) and
                   not lines[i+1].strip().startswith('*End of') and
                   not lines[i+1].strip().startswith('_End of')):
                i += 1
                full_text += ' ' + lines[i].strip()

            add_body_paragraph(doc, full_text, first_line_indent=True)

        i += 1


# -- Main Execution --

def main():
    print("=" * 60)
    print("CRP Appendix Merge - 4 Appendix Structure (A-D)")
    print("=" * 60)

    # Verify all inputs exist
    if not os.path.exists(CRP_PATH):
        print(f"ERROR: CRP document not found: {CRP_PATH}")
        sys.exit(1)
    for letter, filename, title in APPENDIX_FILES:
        filepath = os.path.join(APPENDIX_DIR, filename)
        if not os.path.exists(filepath):
            print(f"ERROR: Appendix {letter} not found: {filepath}")
            sys.exit(1)
        print(f"  Found Appendix {letter}: {filename}")

    print(f"\nOpening CRP document...")
    print(f"  Source: {os.path.basename(CRP_PATH)}")
    print(f"  Size: {os.path.getsize(CRP_PATH):,} bytes")
    doc = Document(CRP_PATH)

    body = doc.element.body
    paragraphs = body.findall(qn('w:p'))
    print(f"  Total paragraphs: {len(paragraphs)}")

    # Find the first APPENDIX heading (Heading 1 style) after the List of Appendices
    # The root CRP has body content through para ~632, then old appendixes start at 633
    appendix_start_idx = None
    list_of_appendices_idx = None

    for idx, p_elem in enumerate(paragraphs):
        text_parts = []
        for t in p_elem.itertext():
            text_parts.append(t)
        full_text = ''.join(text_parts).strip()

        pPr = p_elem.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                if 'List of Appendices' in full_text:
                    list_of_appendices_idx = idx
                elif full_text.upper().startswith('APPENDIX'):
                    if list_of_appendices_idx is not None and appendix_start_idx is None:
                        appendix_start_idx = idx

    if list_of_appendices_idx is None:
        print("ERROR: Could not find 'List of Appendices' heading!")
        sys.exit(1)

    if appendix_start_idx is None:
        print("WARNING: Could not find first APPENDIX heading. Will append after List of Appendices entries.")
        # Fall back: strip from after the List of Appendices + its entries
        # Find the next Heading1 or a reasonable cut point after the list
        appendix_start_idx = list_of_appendices_idx
        # We'll strip from the List of Appendices heading and rebuild it
        strip_includes_list = True
    else:
        strip_includes_list = False

    print(f"\n  List of Appendices at paragraph: {list_of_appendices_idx}")
    print(f"  First APPENDIX heading at paragraph: {appendix_start_idx}")

    # Remove everything from the appendix start point onward
    # (preserving List of Appendices and its entries)
    elements_to_remove = []
    found_start = False
    target_elem = paragraphs[appendix_start_idx]

    for child in list(body):
        if child is target_elem:
            found_start = True
        if found_start and child.tag != qn('w:sectPr'):
            elements_to_remove.append(child)

    print(f"\n  Removing {len(elements_to_remove)} elements (old appendix content)...")
    for elem in elements_to_remove:
        body.remove(elem)

    # Verify removal
    remaining_paras = body.findall(qn('w:p'))
    print(f"  Remaining paragraphs after strip: {len(remaining_paras)}")

    # Process each appendix
    current_landscape = False
    total_elements_added = 0

    for letter, filename, title in APPENDIX_FILES:
        filepath = os.path.join(APPENDIX_DIR, filename)
        print(f"\nProcessing Appendix {letter}: {title}...")
        print(f"  File: {filename}")

        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        print(f"  Lines: {len(lines)}")

        needs_landscape = letter in LANDSCAPE_APPENDIXES

        # If switching orientation, add a section break
        if needs_landscape and not current_landscape:
            add_section_break(doc, landscape=True)
            current_landscape = True
        elif not needs_landscape and current_landscape:
            add_section_break(doc, landscape=False)
            current_landscape = False

        before_count = len(body.findall(qn('w:p')))

        # Add title page
        add_appendix_title_page(doc, letter, title)

        # Process content
        process_markdown_content(doc, lines, landscape=needs_landscape)

        after_count = len(body.findall(qn('w:p')))
        added = after_count - before_count
        total_elements_added += added
        print(f"  Added {added} paragraphs")

    # If we ended in landscape, switch back to portrait
    if current_landscape:
        add_section_break(doc, landscape=False)

    # Update the final section properties
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        pgSz = sect_pr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.set(qn('w:w'), '12240')
            pgSz.set(qn('w:h'), '15840')
            if pgSz.get(qn('w:orient')):
                del pgSz.attrib[qn('w:orient')]

    # Save
    print(f"\n{'=' * 60}")
    print(f"Saving merged document...")
    print(f"  Output: {os.path.basename(OUTPUT_PATH)}")
    doc.save(OUTPUT_PATH)
    output_size = os.path.getsize(OUTPUT_PATH)
    print(f"  Size: {output_size:,} bytes")

    # Summary
    final_paras = len(body.findall(qn('w:p')))
    print(f"\nMerge Summary:")
    print(f"  Original paragraphs: {len(paragraphs)}")
    print(f"  Removed (old appendixes): {len(elements_to_remove)}")
    print(f"  Added (new appendixes): {total_elements_added}")
    print(f"  Final paragraphs: {final_paras}")
    print(f"  Output file: {output_size:,} bytes")
    print(f"\nDone!")


if __name__ == '__main__':
    main()
